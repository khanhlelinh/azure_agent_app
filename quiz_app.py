import sys
import os
import sqlite3
import re
import random
from datetime import datetime
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
                             QPushButton, QFileDialog, QTableWidget, QTableWidgetItem, 
                             QLabel, QHeaderView, QMessageBox, QLineEdit, QSplitter, 
                             QSlider, QStackedWidget, QTextEdit, QScrollArea, QGroupBox, 
                             QCheckBox, QSpinBox, QGridLayout, QComboBox, QRadioButton, QButtonGroup)
from PyQt6.QtCore import Qt, QTimer, QSettings
from PyQt6.QtGui import QColor, QPixmap, QFont

# Thư viện xuất file Word
try:
    from docx import Document
    from docx.shared import RGBColor
except ImportError:
    print("Vui lòng cài đặt thư viện: pip install python-docx")
    sys.exit(1)

# =========================================================
# --- TAILWIND CSS STYLING (HIỆN ĐẠI & TẬP TRUNG) ---
# =========================================================
TAILWIND_QSS = """
    QMainWindow { background-color: #F3F4F6; }
    
    QWidget { font-family: "Segoe UI", sans-serif; color: #111827; font-size: 16px; }
    
    QDialog, QMessageBox {
        background-color: #FFFFFF;
    }
    QMessageBox QLabel {
        color: #111827;
        font-size: 16px;
    }
    
    QGroupBox { 
        background-color: #FFFFFF; 
        border: 1px solid #D1D5DB; 
        border-radius: 10px; 
        margin-top: 15px; 
        padding: 20px 15px 15px 15px;
    }
    QGroupBox::title { 
        subcontrol-origin: margin; 
        left: 15px; 
        padding: 0 8px; 
        color: #4F46E5; 
        font-size: 18px; 
        font-weight: bold;
    }
    
    QLineEdit, QComboBox, QSpinBox {
        border: 1px solid #D1D5DB;
        border-radius: 6px;
        padding: 8px 12px;
        background-color: #FFFFFF;
        color: #111827;
        font-weight: bold;
        font-size: 16px;
    }
    QLineEdit:focus, QComboBox:focus, QSpinBox:focus {
        border: 2px solid #4F46E5;
        background-color: #F9FAFB;
    }
    
    QLineEdit#txtUserAnswer {
        font-size: 18px;
        color: #4F46E5;
        border: 2px solid #9CA3AF;
        background-color: #EEF2FF;
    }
    QLineEdit#txtUserAnswer:focus {
        border-color: #4F46E5;
        background-color: #FFFFFF;
    }
    
    QTextEdit#txtQuestionText {
        background-color: #FFFFFF;
        border: 1px solid #D1D5DB;
        border-radius: 8px;
        padding: 16px;
        color: #111827;
        font-size: 22px; 
        font-weight: 500;
        line-height: 1.4;
    }
    
    QCheckBox#chkExcludeLab {
        font-weight: bold;
        color: #DC2626;
        font-size: 16px;
    }
    
    QRadioButton {
        font-size: 16px;
        font-weight: bold;
        color: #1F2937;
        padding: 5px;
    }
    
    QPushButton {
        background-color: #4F46E5; color: white;
        border-radius: 6px; padding: 10px 16px;
        font-weight: bold; font-size: 16px; border: none;
    }
    QPushButton:hover { background-color: #4338CA; }
    QPushButton:disabled { background-color: #9CA3AF; color: #F3F4F6; }
    
    QPushButton#btnSuccess { background-color: #10B981; }
    QPushButton#btnSuccess:hover { background-color: #059669; }
    
    QPushButton#btnWarning { background-color: #F59E0B; }
    QPushButton#btnWarning:hover { background-color: #D97706; }
    
    QWidget#cheatPanel {
        background-color: #FEF3C7;
        border: 2px solid #F59E0B;
        border-radius: 8px;
        margin-top: 15px;
        margin-bottom: 10px;
    }
    
    QTextEdit#txtCheatExp {
        background-color: transparent; 
        border: none; 
        color: #111827; 
        font-size: 18px; 
    }
    
    QScrollArea { border: none; background-color: transparent; }
    
    /* BẢNG KẾT QUẢ */
    QTableWidget {
        background-color: #FFFFFF;
        alternate-background-color: #F9FAFB;
        color: #1F2937;
        gridline-color: #E5E7EB;
        border: 1px solid #D1D5DB;
        border-radius: 8px;
        font-size: 15px;
    }
    QTableWidget::item {
        padding: 8px;
    }
    QTableWidget::item:selected {
        background-color: #EEF2FF;
        color: #4F46E5;
    }
    
    /* Tách riêng CSS cho cột ngang (Tiêu đề) và cột dọc (STT) */
    QHeaderView::section:horizontal {
        background-color: #F3F4F6;
        color: #374151;
        padding: 10px;
        font-weight: bold;
        font-size: 16px;
        border: none;
        border-right: 1px solid #D1D5DB;
        border-bottom: 2px solid #D1D5DB;
    }
    QHeaderView::section:vertical {
        background-color: #F3F4F6;
        color: #374151;
        padding: 5px; /* Giảm đệm để không chèn chữ số */
        font-weight: bold;
        font-size: 15px;
        border: none;
        border-right: 2px solid #D1D5DB;
        border-bottom: 1px solid #E5E7EB;
    }
"""

def natural_sort_key(s):
    if isinstance(s, tuple):
        s = s[0]
    return [int(text) if text.isdigit() else text.lower() for text in re.split(r'(\d+)', s)]

class QuizAppWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Hệ thống Thiết kế Đề thi Thông minh - Chế độ Động (Topic/Study Area)")
        self.resize(1300, 850)
        self.setStyleSheet(TAILWIND_QSS)
        
        self.settings = QSettings("MyAIWorkspace", "AzureSQLAgent")
        self.base_image_folder = self.settings.value("last_img_folder", "")
        self.db_path = ""
        
        self.questions = []
        self.current_index = 0
        self.user_answers = {}
        self.timer = QTimer()
        self.time_left_seconds = 0
        self.is_submitted = False
        self.is_peeking = False 
        
        self.category_widgets = [] 
        
        self.init_ui()

    def init_ui(self):
        self.main_widget = QWidget()
        self.setCentralWidget(self.main_widget)
        layout = QVBoxLayout(self.main_widget)
        
        self.stacked_widget = QStackedWidget()
        layout.addWidget(self.stacked_widget)
        
        self.page_config = QWidget()
        self.setup_config_page()
        self.stacked_widget.addWidget(self.page_config)
        
        self.page_quiz = QWidget()
        self.setup_quiz_page()
        self.stacked_widget.addWidget(self.page_quiz)
        
        self.page_result = QWidget()
        self.setup_result_page()
        self.stacked_widget.addWidget(self.page_result)

    # =========================================================================
    # --- TRANG 1: BLUEPRINT THIẾT KẾ ĐỀ THI ---
    # =========================================================================
    def setup_config_page(self):
        layout = QVBoxLayout(self.page_config)
        layout.setContentsMargins(40, 20, 40, 20)
        
        top_bar = QHBoxLayout()
        self.btn_load_db = QPushButton("🗄️ Chọn CSDL (.db)")
        self.btn_load_db.clicked.connect(self.open_db)
        self.lbl_db_status = QLabel("Chưa chọn file.")
        self.lbl_db_status.setStyleSheet("font-weight: bold; color: #DC2626;")
        
        self.btn_load_img = QPushButton("📁 Folder Ảnh gốc")
        self.btn_load_img.clicked.connect(self.select_img_folder)
        self.lbl_img_status = QLabel(f"📁: {os.path.basename(self.base_image_folder)}" if self.base_image_folder else "Chưa chọn.")
        self.lbl_img_status.setStyleSheet("font-weight: bold; color: #10B981;")
        
        top_bar.addWidget(self.btn_load_db); top_bar.addWidget(self.lbl_db_status, 1)
        top_bar.addWidget(self.btn_load_img); top_bar.addWidget(self.lbl_img_status, 1)
        layout.addLayout(top_bar)
        
        self.card_blueprint = QGroupBox("📊 Ma trận Phân bổ số lượng câu hỏi")
        card_layout = QVBoxLayout(self.card_blueprint)
        
        mode_layout = QHBoxLayout()
        self.rb_mode_topic = QRadioButton("Lấy câu hỏi theo Topic")
        self.rb_mode_area = QRadioButton("Lấy câu hỏi theo Study Area")
        self.rb_mode_topic.setChecked(True)
        
        self.mode_group = QButtonGroup()
        self.mode_group.addButton(self.rb_mode_topic)
        self.mode_group.addButton(self.rb_mode_area)
        
        self.rb_mode_topic.toggled.connect(self.load_categories_dynamic)
        
        self.chk_exclude_lab = QCheckBox("🚫 Bỏ qua các câu hỏi Lab (Thực hành)")
        self.chk_exclude_lab.setObjectName("chkExcludeLab")
        self.chk_exclude_lab.stateChanged.connect(self.load_categories_dynamic)
        
        mode_layout.addWidget(self.rb_mode_topic)
        mode_layout.addWidget(self.rb_mode_area)
        mode_layout.addStretch()
        mode_layout.addWidget(self.chk_exclude_lab)
        card_layout.addLayout(mode_layout)
        
        self.scroll_cats = QScrollArea()
        self.scroll_cats.setWidgetResizable(True)
        self.scroll_cats.setStyleSheet("background: white; border: 1px solid #E5E7EB; border-radius: 8px;")
        
        self.container_cats = QWidget()
        self.layout_cats_grid = QGridLayout(self.container_cats)
        self.layout_cats_grid.setAlignment(Qt.AlignmentFlag.AlignTop)
        self.layout_cats_grid.setHorizontalSpacing(30) 
        self.layout_cats_grid.setVerticalSpacing(15)   
        
        self.scroll_cats.setWidget(self.container_cats)
        card_layout.addWidget(self.scroll_cats)
        
        footer_layout = QHBoxLayout()
        
        info_vbox = QVBoxLayout()
        self.lbl_total_selected = QLabel("Tổng số câu hỏi sẽ thi: 0")
        self.lbl_total_selected.setStyleSheet("font-size: 20px; font-weight: bold; color: #4F46E5;")
        
        time_config_layout = QHBoxLayout()
        time_config_layout.addWidget(QLabel("⏳ Số giây cho mỗi câu hỏi:"))
        self.spin_seconds_per_q = QSpinBox()
        self.spin_seconds_per_q.setRange(10, 600) 
        self.spin_seconds_per_q.setValue(60) 
        self.spin_seconds_per_q.setFixedWidth(80)
        time_config_layout.addWidget(self.spin_seconds_per_q)
        time_config_layout.addStretch()
        
        info_vbox.addWidget(self.lbl_total_selected)
        info_vbox.addLayout(time_config_layout)
        
        self.btn_start = QPushButton("🚀 BẮT ĐẦU THI")
        self.btn_start.setObjectName("btnSuccess")
        self.btn_start.setMinimumSize(250, 60)
        self.btn_start.setEnabled(False)
        self.btn_start.clicked.connect(self.start_exam)
        
        footer_layout.addLayout(info_vbox)
        footer_layout.addStretch()
        footer_layout.addWidget(self.btn_start)
        card_layout.addLayout(footer_layout)
        
        layout.addWidget(self.card_blueprint)

    def open_db(self):
        path, _ = QFileDialog.getOpenFileName(self, "Chọn CSDL SQLite", "", "SQLite Database (*.db *.sqlite)")
        if path:
            self.db_path = path
            self.lbl_db_status.setText(f"🗄️: {os.path.basename(path)}")
            self.lbl_db_status.setStyleSheet("font-weight: bold; color: #10B981;")
            self.load_categories_dynamic()

    def select_img_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Chọn thư mục ảnh", self.base_image_folder)
        if folder:
            self.base_image_folder = folder
            self.settings.setValue("last_img_folder", folder)
            self.lbl_img_status.setText(f"📁: {os.path.basename(folder)}")

    def load_categories_dynamic(self):
        for i in reversed(range(self.layout_cats_grid.count())):
            w = self.layout_cats_grid.itemAt(i).widget()
            if w: w.deleteLater()
        self.category_widgets = []

        if not self.db_path: return

        try:
            is_exclude_lab = self.chk_exclude_lab.isChecked()
            target_column = "category" if self.rb_mode_topic.isChecked() else "study_area"
            
            sql = f"SELECT {target_column}, COUNT(*) FROM questions WHERE {target_column} IS NOT NULL AND trim({target_column}) != ''"
            if is_exclude_lab:
                sql += " AND question_type != 'Lab'" 
            sql += f" GROUP BY {target_column}"

            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            try:
                cursor.execute(sql)
                raw_data = cursor.fetchall()
            except sqlite3.OperationalError:
                QMessageBox.warning(self, "CSDL Chưa được nâng cấp", "Vui lòng mở file DB này bằng module DB Manager trước để hệ thống tự động sinh cột 'Study Area'.")
                conn.close()
                return
                
            conn.close()

            sorted_data = sorted(raw_data, key=natural_sort_key)

            for idx, (cat_name, count) in enumerate(sorted_data):
                if count <= 0: continue
                
                topic_card = QWidget()
                topic_layout = QHBoxLayout(topic_card)
                topic_layout.setContentsMargins(5, 2, 5, 2)
                topic_layout.setSpacing(8)
                
                name_lbl = QLabel(f"<b>{cat_name}</b>")
                name_lbl.setFixedWidth(180) 
                name_lbl.setToolTip(cat_name) 
                
                slider = QSlider(Qt.Orientation.Horizontal)
                slider.setRange(0, count)
                slider.setValue(0)
                slider.setFixedWidth(140) 
                
                spin = QSpinBox()
                spin.setRange(0, count)
                spin.setValue(0)
                spin.setFixedWidth(90) 
                
                max_lbl = QLabel(f"/ {count}")
                max_lbl.setStyleSheet("color: #6B7280; font-size: 14px;")
                max_lbl.setFixedWidth(70)

                slider.valueChanged.connect(spin.setValue)
                spin.valueChanged.connect(slider.setValue)
                slider.valueChanged.connect(self.calculate_total)

                topic_layout.addWidget(name_lbl)
                topic_layout.addWidget(slider)
                topic_layout.addWidget(spin)
                topic_layout.addWidget(max_lbl)
                
                curr_idx = len(self.category_widgets)
                row = curr_idx // 2
                col = curr_idx % 2
                self.layout_cats_grid.addWidget(topic_card, row, col)
                
                self.category_widgets.append({
                    'name': cat_name,
                    'slider': slider,
                    'spin': spin
                })
            
            self.calculate_total()
        except Exception as e:
            QMessageBox.critical(self, "Lỗi nạp danh mục", str(e))

    def calculate_total(self):
        total = sum(item['slider'].value() for item in self.category_widgets)
        self.lbl_total_selected.setText(f"Tổng số câu hỏi sẽ thi: {total}")
        self.btn_start.setEnabled(total > 0)

    def start_exam(self):
        if not self.build_exam_questions(): return
        
        self.current_index = 0
        self.user_answers = {}
        self.is_submitted = False
        self.is_peeking = False
        self.wrong_questions = []
        
        self.render_question()
        
        total_questions = len(self.questions)
        secs_per_q = self.spin_seconds_per_q.value()
        self.time_left_seconds = total_questions * secs_per_q
        
        mins, secs = divmod(self.time_left_seconds, 60)
        self.lbl_timer.setText(f"⏱️ {mins:02d}:{secs:02d}")
        self.lbl_timer.setStyleSheet("font-size: 20px; font-weight: bold; color: #DC2626;")
        
        self.timer = QTimer(self)
        self.timer.timeout.connect(self.update_timer)
        self.timer.start(1000)
        
        self.stacked_widget.setCurrentWidget(self.page_quiz)

    def build_exam_questions(self):
        conn = None
        try:
            is_exclude_lab = self.chk_exclude_lab.isChecked()
            target_column = "category" if self.rb_mode_topic.isChecked() else "study_area"
            
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            all_selected = []

            for item in self.category_widgets:
                qty = item['slider'].value()
                if qty > 0:
                    sql = f"SELECT source_name, extracted_text, choices, final_answer, answer, question_type FROM questions WHERE {target_column} = ?"
                    params = [item['name']]
                    
                    if is_exclude_lab:
                        sql += " AND question_type != 'Lab'"
                        
                    sql += " ORDER BY RANDOM() LIMIT ?"
                    params.append(qty)
                    
                    cursor.execute(sql, params)
                    all_selected.extend(cursor.fetchall())
            
            if not all_selected: return False

            random.shuffle(all_selected)
            
            self.questions = []
            for r in all_selected:
                self.questions.append({
                    'source': r[0], 'text': r[1] or "", 'choices': r[2] or "",
                    'correct_ans': (r[3] or "").strip().upper(),
                    'exp': r[4] or "", 'type': r[5] or "Chọn"
                })
            return True
        except Exception as e:
            QMessageBox.critical(self, "Lỗi tạo đề", str(e))
            return False
        finally:
            if conn: conn.close()

    # =========================================================================
    # --- TRANG 2: KHU VỰC THI ---
    # =========================================================================
    def setup_quiz_page(self):
        layout = QVBoxLayout(self.page_quiz)
        layout.setContentsMargins(60, 20, 60, 20)
        
        header_layout = QHBoxLayout()
        self.lbl_q_progress = QLabel("Câu hỏi: 1 / 50")
        self.lbl_q_progress.setStyleSheet("font-size: 20px; font-weight: bold; color: #4F46E5;")
        self.lbl_timer = QLabel("⏱️ 00:00")
        self.lbl_timer.setStyleSheet("font-size: 20px; font-weight: bold; color: #DC2626;")
        
        header_layout.addWidget(self.lbl_q_progress)
        header_layout.addStretch()
        header_layout.addWidget(self.lbl_timer)
        layout.addLayout(header_layout)
        
        self.txt_question_text = QTextEdit()
        self.txt_question_text.setObjectName("txtQuestionText")
        self.txt_question_text.setReadOnly(True)
        self.txt_question_text.setMinimumHeight(180) 
        layout.addWidget(self.txt_question_text, stretch=2) 
        
        self.master_scroll = QScrollArea()
        self.master_scroll.setWidgetResizable(True)
        self.master_scroll.setStyleSheet("border: none; background: transparent;")
        
        self.scroll_content = QWidget()
        self.scroll_layout = QVBoxLayout(self.scroll_content)
        self.scroll_layout.setAlignment(Qt.AlignmentFlag.AlignTop)
        self.scroll_layout.setContentsMargins(0, 0, 0, 0)
        
        self.container_q_img = QWidget()
        self.layout_q_img = QVBoxLayout(self.container_q_img)
        self.layout_q_img.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.layout_q_img.setContentsMargins(0, 0, 0, 0)
        self.scroll_layout.addWidget(self.container_q_img)
        
        self.cheat_panel = QWidget()
        self.cheat_panel.setObjectName("cheatPanel")
        cheat_layout = QVBoxLayout(self.cheat_panel)
        cheat_layout.setContentsMargins(20, 16, 20, 16)
        
        self.lbl_cheat_correct_ans = QLabel()
        self.lbl_cheat_correct_ans.setStyleSheet("color: #D97706; font-size: 20px; font-weight: bold;")
        cheat_layout.addWidget(self.lbl_cheat_correct_ans)
        
        self.txt_cheat_exp = QTextEdit()
        self.txt_cheat_exp.setObjectName("txtCheatExp")
        self.txt_cheat_exp.setReadOnly(True)
        self.txt_cheat_exp.setMinimumHeight(150) 
        cheat_layout.addWidget(self.txt_cheat_exp)
        
        self.cheat_img_container = QWidget()
        self.cheat_img_layout = QVBoxLayout(self.cheat_img_container)
        self.cheat_img_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.cheat_img_layout.setContentsMargins(0, 0, 0, 0)
        cheat_layout.addWidget(self.cheat_img_container)
        
        self.cheat_panel.setVisible(False) 
        self.scroll_layout.addWidget(self.cheat_panel)
        
        self.master_scroll.setWidget(self.scroll_content)
        layout.addWidget(self.master_scroll, stretch=3) 
        
        ans_input_layout = QHBoxLayout()
        ans_input_layout.setContentsMargins(0, 10, 0, 10)
        ans_input_layout.addWidget(QLabel("<b>👉 Gõ đáp án của bạn:</b>"))
        
        self.txt_user_ans = QLineEdit()
        self.txt_user_ans.setObjectName("txtUserAnswer")
        self.txt_user_ans.setPlaceholderText("Gõ A, B hoặc A,C,D,E...")
        self.txt_user_ans.textChanged.connect(self.save_current_answer)
        ans_input_layout.addWidget(self.txt_user_ans, stretch=1)
        layout.addLayout(ans_input_layout)
        
        nav_layout = QHBoxLayout()
        self.btn_quiz_prev = QPushButton("◀ Câu trước")
        self.btn_quiz_prev.clicked.connect(self.prev_question)
        
        self.btn_peek = QPushButton("👀 Nhìn trộm đáp án")
        self.btn_peek.setObjectName("btnWarning")
        self.btn_peek.clicked.connect(self.toggle_peeking)
        
        self.btn_quiz_next = QPushButton("Câu sau ▶")
        self.btn_quiz_next.clicked.connect(self.next_question)
        
        self.btn_submit = QPushButton("📝 NỘP BÀI")
        self.btn_submit.setObjectName("btnSuccess")
        self.btn_submit.clicked.connect(self.confirm_submit)
        
        nav_layout.addWidget(self.btn_quiz_prev)
        nav_layout.addWidget(self.btn_peek)
        nav_layout.addStretch()
        nav_layout.addWidget(self.btn_quiz_next)
        nav_layout.addWidget(self.btn_submit)
        layout.addLayout(nav_layout)

    def update_timer(self):
        if self.time_left_seconds <= 0:
            self.timer.stop()
            self.lbl_timer.setText("⏱️ HẾT GIỜ!")
            QMessageBox.warning(self, "Hết thời gian", "Đã hết thời gian làm bài. Hệ thống sẽ tự động nộp bài!")
            self.finish_and_grade_quiz()
            return
            
        self.time_left_seconds -= 1
        mins, secs = divmod(self.time_left_seconds, 60)
        self.lbl_timer.setText(f"⏱️ {mins:02d}:{secs:02d}")
        
        if self.time_left_seconds <= 10:
            if self.time_left_seconds % 2 == 0:
                self.lbl_timer.setStyleSheet("font-size: 20px; font-weight: bold; color: #DC2626;")
            else:
                self.lbl_timer.setStyleSheet("font-size: 20px; font-weight: bold; color: #111827;")

    def render_question(self):
        if not self.questions: return
        q = self.questions[self.current_index]
        
        self.is_peeking = False 
        self.cheat_panel.setVisible(False)
        self.btn_peek.setText("👀 Nhìn trộm đáp án")
        
        self.lbl_q_progress.setText(f"Câu hỏi: {self.current_index + 1} / {len(self.questions)}")
        
        combined_text = q['text']
        if q['choices'].strip():
            combined_text += "\n\n" + q['choices']
        self.txt_question_text.setPlainText(combined_text.strip())
        
        saved_text = self.user_answers.get(self.current_index, "")
        self.txt_user_ans.blockSignals(True)
        self.txt_user_ans.setText(saved_text)
        self.txt_user_ans.blockSignals(False)
        
        for i in reversed(range(self.layout_q_img.count())):
            w = self.layout_q_img.itemAt(i).widget()
            if w: w.deleteLater()
            
        for i in reversed(range(self.cheat_img_layout.count())):
            w = self.cheat_img_layout.itemAt(i).widget()
            if w: w.deleteLater()
            
        ans_images = [] 

        if self.base_image_folder:
            full_dir = os.path.join(self.base_image_folder, q['source'].replace("📁 ", "").strip())
            if os.path.isdir(full_dir):
                files = [f for f in os.listdir(full_dir) if f.lower().endswith(('.png', '.jpg', '.jpeg'))]
                for fname in sorted(files, key=natural_sort_key):
                    img_path = os.path.join(full_dir, fname)
                    if "-answer" in fname.lower():
                        ans_images.append(img_path)
                    else:
                        lbl = QLabel(); lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
                        pix = QPixmap(img_path)
                        if not pix.isNull():
                            lbl.setPixmap(pix.scaled(1000, 2000, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation))
                            self.layout_q_img.addWidget(lbl)

        self.lbl_cheat_correct_ans.setText(f"🎯 Đáp án CHUẨN: {q['correct_ans'] if q['correct_ans'] else 'CHƯA CÓ TRONG DB'}")
        self.txt_cheat_exp.setPlainText(q['exp'] if q['exp'].strip() else "Không có lời giải thích dạng văn bản.")
        for p in ans_images:
            lbl = QLabel(); lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
            pix = QPixmap(p)
            if not pix.isNull():
                lbl.setPixmap(pix.scaled(900, 2000, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation))
                self.cheat_img_layout.addWidget(lbl)
        
        self.btn_quiz_prev.setEnabled(self.current_index > 0)
        self.btn_quiz_next.setEnabled(self.current_index < len(self.questions)-1)
        
        self.master_scroll.verticalScrollBar().setValue(0)
        self.txt_user_ans.setFocus()

    def save_current_answer(self):
        if self.is_submitted: return
        raw_text = self.txt_user_ans.text()
        clean_parts = [p.strip().upper() for p in raw_text.split(",") if p.strip()]
        self.user_answers[self.current_index] = ",".join(clean_parts)

    def toggle_peeking(self):
        self.is_peeking = not self.is_peeking
        self.cheat_panel.setVisible(self.is_peeking)
        
        if self.is_peeking:
            self.btn_peek.setText("🙈 Giấu đáp án đi")
            QTimer.singleShot(100, lambda: self.master_scroll.verticalScrollBar().setValue(
                self.master_scroll.verticalScrollBar().maximum()
            ))
        else:
            self.btn_peek.setText("👀 Nhìn trộm đáp án")

    def goto_question(self, index):
        self.current_index = index; self.render_question()

    def prev_question(self):
        if self.current_index > 0: self.goto_question(self.current_index - 1)

    def next_question(self):
        if self.current_index < len(self.questions)-1: self.goto_question(self.current_index + 1)

    def confirm_submit(self):
        missing = sum(1 for i in range(len(self.questions)) if not self.user_answers.get(i, "").strip())
        msg = f"Nộp bài thi? {f'(Còn {missing} câu chưa điền đáp án)' if missing > 0 else ''}"
        if QMessageBox.question(self, "Xác nhận", msg, QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No) == QMessageBox.StandardButton.Yes:
            self.finish_and_grade_quiz()

    # =========================================================================
    # --- TRANG 3: KẾT QUẢ VÀ BÁO CÁO (BẢNG SÁNG TRẮNG CHUẨN UI) ---
    # =========================================================================
    def finish_and_grade_quiz(self):
        self.is_submitted = True; self.timer.stop()
        correct_cnt = 0
        self.wrong_questions = []
        
        for idx, q in enumerate(self.questions):
            u_ans = self.user_answers.get(idx, "").strip()
            c_ans = q['correct_ans']
            
            if u_ans == c_ans and c_ans != "":
                correct_cnt += 1
            else:
                self.wrong_questions.append({
                    'source': q['source'], 
                    'text': q['text'], 
                    'user_ans': u_ans if u_ans else "BỎ TRỐNG", 
                    'correct_ans': c_ans if c_ans else "TRỐNG", 
                    'exp': q['exp']
                })
                
        self.lbl_score.setText(f"🎯 Điểm: {(correct_cnt/len(self.questions)*100):.1f} / 100")
        self.lbl_correct_count.setText(f"✅ Đúng: {correct_cnt} / {len(self.questions)}")
        
        total_initial_time = len(self.questions) * self.spin_seconds_per_q.value()
        time_used = total_initial_time - self.time_left_seconds
        mins_used, secs_used = divmod(time_used, 60)
        self.lbl_time_taken.setText(f"⏱️ Thời gian dùng: {mins_used:02d}:{secs_used:02d}")
        
        self.table_wrong.setAlternatingRowColors(True) 
        self.table_wrong.setRowCount(0)
        
        # Thiết lập độ rộng cột STT luôn luôn vừa vặn không bị đè
        self.table_wrong.verticalHeader().setMinimumWidth(50)
        
        for w in self.wrong_questions:
            r = self.table_wrong.rowCount(); self.table_wrong.insertRow(r)
            
            item_source = QTableWidgetItem(w['source'])
            self.table_wrong.setItem(r, 0, item_source)
            
            item_user = QTableWidgetItem(w['user_ans'])
            item_user.setForeground(QColor("#DC2626")) 
            item_user.setFont(QFont("Segoe UI", 12, QFont.Weight.Bold))
            item_user.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.table_wrong.setItem(r, 1, item_user)
            
            item_correct = QTableWidgetItem(w['correct_ans'])
            item_correct.setForeground(QColor("#10B981")) 
            item_correct.setFont(QFont("Segoe UI", 12, QFont.Weight.Bold))
            item_correct.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.table_wrong.setItem(r, 2, item_correct)
            
        self.stacked_widget.setCurrentWidget(self.page_result)

    def setup_result_page(self):
        layout = QVBoxLayout(self.page_result)
        summary_group = QGroupBox("📊 Kết quả")
        summary_layout = QHBoxLayout(summary_group)
        self.lbl_score = QLabel(); self.lbl_correct_count = QLabel(); self.lbl_time_taken = QLabel()
        summary_layout.addWidget(self.lbl_score); summary_layout.addWidget(self.lbl_correct_count); summary_layout.addWidget(self.lbl_time_taken)
        layout.addWidget(summary_group)
        
        self.table_wrong = QTableWidget(0, 3)
        self.table_wrong.setHorizontalHeaderLabels(["Nguồn (Source)", "Bạn gõ", "Đáp án chuẩn"])
        self.table_wrong.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        self.table_wrong.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.Fixed)
        self.table_wrong.setColumnWidth(1, 150)
        self.table_wrong.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.Fixed)
        self.table_wrong.setColumnWidth(2, 150)
        layout.addWidget(self.table_wrong, 1)
        
        btn_layout = QHBoxLayout()
        self.btn_export_word = QPushButton("📄 Xuất Word"); self.btn_export_word.clicked.connect(self.export_wrong_to_word)
        self.btn_restart = QPushButton("🔄 Thi lại"); self.btn_restart.clicked.connect(lambda: self.stacked_widget.setCurrentWidget(self.page_config))
        btn_layout.addStretch(); btn_layout.addWidget(self.btn_export_word); btn_layout.addWidget(self.btn_restart)
        layout.addLayout(btn_layout)

    def export_wrong_to_word(self):
        path, _ = QFileDialog.getSaveFileName(self, "Lưu báo cáo", "Bao_Cao_Cau_Sai.docx", "Word (*.docx)")
        if path:
            doc = Document(); doc.add_heading("BÁO CÁO ÔN TẬP", 0)
            for idx, w in enumerate(self.wrong_questions):
                doc.add_heading(f"Câu {idx+1}: {w['source']}", level=2)
                doc.add_paragraph(w['text'])
                p = doc.add_paragraph(); p.add_run("❌ Bạn gõ: ").bold = True; p.add_run(w['user_ans']).font.color.rgb = RGBColor(255, 0, 0)
                p2 = doc.add_paragraph(); p2.add_run("✅ Chuẩn: ").bold = True; p2.add_run(w['correct_ans']).font.color.rgb = RGBColor(0, 128, 0)
                doc.add_paragraph(f"💡 Giải thích: {w['exp']}")
            doc.save(path)

if __name__ == '__main__':
    app = QApplication(sys.argv); window = QuizAppWindow(); window.show(); sys.exit(app.exec())